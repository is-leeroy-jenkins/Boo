'''
  ******************************************************************************************
      Assembly:                Boo
      Filename:                Tigrr.py
      Author:                  Terry D. Eppler
      Created:                 05-31-2022

      Last Modified By:        Terry D. Eppler
      Last Modified On:        05-01-2025
  ******************************************************************************************
  <copyright file="tigrr.py" company="Terry D. Eppler">

	     Boo is a df analysis tool that integrates various Generative AI, Text-Processing, and
	     Machine-Learning algorithms for federal analysts.
	     Copyright ©  2022  Terry Eppler

     Permission is hereby granted, free of charge, to any person obtaining a copy
     of this software and associated documentation files (the “Software”),
     to deal in the Software without restriction,
     including without limitation the rights to use,
     copy, modify, merge, publish, distribute, sublicense,
     and/or sell copies of the Software,
     and to permit persons to whom the Software is furnished to do so,
     subject to the following conditions:

     The above copyright notice and this permission notice shall be included in all
     copies or substantial portions of the Software.

     THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED,
     INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
     FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT.
     IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
     DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
     ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER
     DEALINGS IN THE SOFTWARE.

     You can contact me at:  terryeppler@gmail.com or eppler.terry@epa.gov

  </copyright>
  <summary>
    Tigrr.py
  </summary>
  ******************************************************************************************
  '''
import os
from collections import defaultdict
from docx import Document as Docx
import re
import json
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup
from openai import OpenAI

from boogr import Error, ErrorDialog
from pathlib import Path
import nltk
from nltk import pos_tag, FreqDist, ConditionalFreqDist
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer, PorterStemmer
from collections import Counter, defaultdict
from docx import Document
import string
import spacy
from scipy import spatial
from sklearn.decomposition import PCA
from sklearn.manifold import TSNE
from sklearn.metrics import average_precision_score, precision_recall_curve
from sklearn.feature_extraction.text import TfidfVectorizer
from pydantic import BaseModel, Field, validator
from pymupdf import Page, Document
import tiktoken
from tiktoken.core import  Encoding
from gensim.models import Word2Vec
from transformers import AutoTokenizer, PreTrainedTokenizerBase
import textwrap as tr
from typing import Any, List, Tuple, Optional, Union, Dict
import unicodedata


# Ensure punkt tokenizer is available for sentence splitting
try:
	nltk.data.find( 'tokenizers/punkt' )
except LookupError:
	nltk.download( 'punkt' )
	nltk.download( 'punkt_tab' )
	nltk.download( 'stopwords' )
	nltk.download( 'wordnet' )
	nltk.download( 'omw-1.4' )
	nltk.download( 'words' )


class Text( BaseModel ):
	'''
	
		Purpose:
		---------
		Class providing path preprocessing functionality
		
	    Methods:
	    --------
	    load_text( url: str ) -> str
	    split_lines( self, path: str ) -> list
	    split_pages( self, path: str, delimit: str ) -> list
	    collapse_whitespace( self, path: str ) -> str
	    remove_punctuation( self, path: str ) -> str:
		remove_special( self, path: str, keep_spaces: bool ) -> str:
		remove_html( self, path: str ) -> str
		remove_errors( self, path: str ) -> str
		correct_errors( self, path: str ) -> str:
		remove_markdown( self, path: str ) -> str
		remove_stopwords( self, path: str ) -> str
		remove_headers( self, pages, min: int=3 ) -> str
	    normalize_text( path: str ) -> str
	    lemmatize_tokens( tokens: List[ str ] ) -> str
	    tokenize_text( path: str ) -> str
	    tokenize_words( path: str ) -> List[ str ]
	    tokenize_sentences( path: str ) -> str
	    chunk_text( self, path: str, max: int=800 ) -> List[ str ]
	    chunk_words( self, path: str, max: int=800, over: int=50 ) -> List[ str ]
	    split_paragraphs( self, path: str ) -> List[ str ]
	    compute_frequency_distribution( self, lines: List[ str ], proc: bool=True ) -> List[ str ]
	    compute_conditional_distribution( self, lines: List[ str ], condition: str=None,
	    proc: bool=True ) -> List[ str ]
	    create_vocabulary( self, frequency, min: int=1 ) -> List[ str ]
	    create_wordbag( tokens: List[ str ] ) -> dict
	    create_word2vec( sentences: List[ str ], vector_size=100, window=5, min_count=1 ) ->
	    Word2Vec
	    create_tfidf( tokens: List[ str ], max_features=1000, prep=True ) -> tuple
	    
	'''
	lemmatizer: WordNetLemmatizer
	stemmer: PorterStemmer
	words: Optional[ List[ str ] ]
	tokens: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	pages: Optional[ List[ str ] ]
	paragraphs: Optional[ List[ str ] ]
	chunck: Optional[ List[ str ] ]
	cleaned_lines: Optional[ List[ str ] ]
	cleaned_tokens: Optional[ List[ str ] ]
	cleaned_pages: Optional[ List[ str ] ]
	removed: Optional[ List[ str ] ]
	raw_pages: Optional[ List[ str ] ]
	stop_words: Optional[ List[ str ] ]
	frequency_distribution: Optional[ Dict[ str, float ] ]
	conditional_distribution: Optional[ Dict[ str, float ] ]


	def __init__( self ):
		'''

			Purpose:
			---------
			Constructor for 'Text' objects

		'''
		super( ).__init__( )
		self.lemmatizer = WordNetLemmatizer( )
		self.stemmer = PorterStemmer( )
		self.words = [ ]
		self.tokens = [ ]
		self.lines = [ ]
		self.pages = [ ]
		self.ids = [ ]
		self.paragraphs = [ ]
		self.chunks = [ ]
		self.chunk_size = 0
		self.cleaned_lines = [ ]
		self.cleaned_tokens = [ ]
		self.cleaned_pages = [ ]
		self.removed = List[ str ]
		self.raw_pages = List[ str ]
		self.stop_words = List[ str ]
		self.frequency_distribution = { }
		self.conditional_distribution = { }
		self.encoding = None
		self.vocabulary = None
		self.file_path = None
		self.raw_input = None
		self.normalized = None
		self.lemmatized = None
		self.tokenized = None
		self.corrected = None
		self.cleaned_text = None
		self.lowercase = None
		self.translator = None
		self.tokenizer = None
		self.vectorizer = None
	
	
	def __dir__( self ) -> List[ str ] | None:
		'''
		
			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None
			
		'''
		return [ 'file_path', 'raw_input', 'raw_pages', 'normalized', 'lemmatized',
		         'tokenized', 'corrected', 'cleaned_text', 'words', 'paragraphs',
		         'tokens', 'pages', 'chunks', 'chunk_size', 'cleaned_pages',
		         'stop_words', 'cleaned_lines', 'removed', 'lowercase', 'encoding', 'vocabulary',
		         'translator', 'lemmatizer', 'stemmer', 'tokenizer', 'vectorizer',
		         'load_text', 'split_lines', 'split_pages', 'collapse_whitespace',
		         'remove_punctuation', 'remove_special', 'remove_html', 'remove_errors',
		         'remove_markdown', 'remove_stopwords', 'remove_headers', 'tiktokenize'
		         'normalize_text', 'lemmatize', 'tokenize_text', 'tokenize_words',
		         'tokenize_sentences', 'chunk_text', 'chunk_words',
		         'create_wordbag', 'create_word2vec', 'create_tfidf',
		         'clean_files', 'convert_jsonl', 'conditional_distribution' ]
	
	
	def load_text( self, path: str ) -> str:
		try:
			if path is None:
				raise Exception( 'The argument "path" is required' )
			else:
				self.file_path = path
				self.raw_input = Path( self.file_path ).read_text( encoding='utf-8' )
				return self.raw_input
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'load_text( self, path: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def collapse_whitespace( self, text: str ) -> str:
		"""

			Purpose:
			-----------
			Removes extra spaces and blank tokens from the path path.

			Parameters:
			-----------
			- path : str

			Returns:
			--------
			A cleaned_lines path path with:
				- Consecutive whitespace reduced to a single space
				- Leading/trailing spaces removed
				- Blank tokens removed
					
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.words = re.sub( r'[ \t]+', ' ', self.raw_input )
				self.cleaned_lines = [ line.strip( ) for line in self.words.splitlines( ) ]
				self.lines = [ line for line in self.cleaned_lines if line ]
				return ' '.join( self.lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'collapse_whitespace( self, path: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def correct_errors( self, text: str ) -> str:
		"""

			Purpose:
			-----------
			Corrects misspelled words in the path path path.

			This function:
			  - Converts path to lowercase
			  - Tokenizes the path into words
			  - Applies spelling correction using TextBlob
			  - Reconstructs and returns the corrected path

			Parameters:
			-----------
			- pages : str
				The path pages path with potential spelling mistakes.

			Returns:
			--------
			- str
				A corrected version of the path path with proper English words.

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.lowercase = self.raw_input.lower( )
				self.tokens = nltk.word_tokenize( self.lowercase )
				self.words = [ str( Word( w ).correct( ) ) for w in self.tokens ]
				_retval = ' '.join( self.words )
				return _retval
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'correct_errors( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_punctuation( self, text: str ) -> str:
		"""

			Purpose:
			-----------
			Removes all punctuation characters from the path path path.

			Parameters:
			-----------
			- pages : str
				The path path path to be cleaned_lines.

			Returns:
			--------
			- str
				The path path with all punctuation removed.

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.translator = str.maketrans( '', '', string.punctuation )
				self.cleaned_text = self.raw_input.translate( self.translator )
				return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_punctuation( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_special( self, text: str ) -> str:
		"""

			Purpose:
			-----------
			Removes special characters from the path path path.

			This function:
			  - Retains only alphanumeric characters and whitespace
			  - Removes symbols like @, #, $, %, &, etc.
			  - Preserves letters, numbers, and spaces

			Parameters:
			-----------
			- pages : str
				The raw path path path potentially
				containing special characters.

			Returns:
			--------
			- str
				A cleaned_lines path containing
				only letters, numbers, and spaces.

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				cleaned = [ ]
				keepers = [ '(', ')', '$', '.', ';', ':' ]
				for char in text:
					if char in keepers:
						cleaned.append( char )
					elif char.isalnum( ) or char.isspace( ):
						cleaned.append( char )
				return ''.join( cleaned )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_special( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_html( self, text: str ) -> str:
		"""

			Purpose:
			--------
			Removes HTML tags from the path path path.
	
			This function:
			  - Parses the path as HTML
			  - Extracts and returns only the visible content without tags
	
			Parameters:
			-----------
			- pages : str
				The path path containing HTML tags.
	
			Returns:
			--------
			- str
				A cleaned_lines path with all HTML tags removed.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_html = text
				self.cleaned_html = BeautifulSoup( self.raw_html, 'raw_html.parser' )
				_retval = self.cleaned_html.get_text( separator=' ', strip=True )
				return _retval
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_html( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_errors( self, text: str ) -> str:
		"""
			
			Purpose:
			_______
			Remove non-English/misspelled words, but preserve numbers and selected symbols.
			Preserved symbols: ( ) $

			Parameters:
			-----------
			- text (str): Input text string.
		
			Returns:
			--------
			- str: Cleaned text.
				
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.lowercase = self.raw_input.lower( )
				self.vocabulary = set( w.lower( ) for w in words.words( ) )
				allowed_symbols = { '(', ')', '$' }
				self.tokens = re.findall( r'\b[\w.]+\b|[()$]', self.lowercase )
				def is_valid_token( token: str ) -> bool:
					return (token in self.vocabulary or token.replace( '.', '',
						1 ).isdigit( ) or token in allowed_symbols)
			self.cleaned_lines = [ tok for tok in self.tokens if is_valid_token( tok ) ]
			return ' '.join( self.cleaned_lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_errors( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )


	def remove_html( self, text: str ) -> str:
		"""
	
			
			Purpose:
			-----------
			 Removes HTML from pages.
	
			Parameters:
			-----------
			- pages : str
				The formatted path pages.
	
			Returns:
			--------
			- str
				A cleaned_lines version of the pages with formatting removed.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				_retval = (BeautifulSoup( self.raw_input, 'raw_html.parser' )
				           .get_text( separator=' ', strip=True ))
				return _retval
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_html( self, path: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )

	
	def remove_markdown( self, text: str ) -> str:
		"""
	
	
			Purpose:
			-----------
			Removes Markdown syntax (e.g., *, #, [], etc.)
	
			Parameters:
			-----------
			- pages : str
				The formatted path pages.
	
			Returns:
			--------
			- str
				A cleaned_lines version of the pages with formatting removed.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.cleaned_text = re.sub( r'\[.*?\]\(.*?\)', '', text )
				self.corrected = re.sub( r'[`_*#~>-]', '', self.cleaned_text )
				_retval = re.sub( r'!\[.*?\]\(.*?\)', '', self.corrected )
				return _retval
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Text'
			exception.method = 'remove_markdown( self, path: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_stopwords( self, text: str ) -> str:
		"""

			Purpose:
			-----------
			This function:
			  - Removes English stopwords from the path pages path.
			  - Tokenizes the path pages
			  - Removes common stopwords (e.g., "the", "is", "and", etc.)
			  - Returns the pages with only meaningful words


			Parameters:
			-----------
			- pages : str
				The text string.
	
			Returns:
			--------
			- str
				A text string without stopwords.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text.lower( )
				self.stop_words = set( stopwords.words( 'english' ) )
				self.tokens = nltk.word_tokenize( text )
				self.cleaned_tokens = [ w for w in self.tokens
				                        if w.isalnum( ) and w not in self.stop_words ]
				self.cleaned_text = ' '.join( self.cleaned_tokens )
				return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Text'
			exception.method = 'remove_stopwords( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
		
		
	def clean_space( self, text: str ) -> str:
		"""

			Purpose:
			_____________
	        Removes extra spaces and blank tokens from the path pages.

	        Parameters:
	        -----------
	        - pages : str
	            The raw path pages path to be cleaned_lines.

	        Returns:
	        --------
	        - str
	            A cleaned_lines pages path with:
	                - Consecutive whitespace reduced to a single space
	                - Leading/trailing spaces removed
	                - Blank tokens removed

	    """
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text.lower( )
				tabs = re.sub( r'[ \t+]', ' ', text.lower( ) )
				collapsed = re.sub( r'\s+', ' ', tabs )
				self.cleaned_lines = [ line for line in collapsed if line ]
				self.cleaned_text = ''.join( cleaned_lines )
				return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Text'
			exception.method = 'remove_errors( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def remove_headers( self, pages: List[ str ], min: int=3 ) -> List[ str ]:
		"""

			Purpose:
			--------
			Removes repetitive headers and footers across a list of pages by frequency analysis.

			Parameters:
			-----------
			- pages (list of str):
				A list where each element is the full path of one page.

			- min (int):
				Minimum num of times a line must appear at the top/bottom to be a header/footer.
		
			Returns:
			---------
			- list of str:
				List of cleaned_lines page tokens without detected headers/footers.
			
		"""
		try:
			if pages is None:
				raise Exception( 'The argument "pages" is required.' )
			else:
				_headers = defaultdict( int )
				_footers = defaultdict( int )
				self.pages = pages
				
				# First pass: collect frequency of top/bottom tokens
				for _page in self.pages:
					self.lines = _page.strip( ).splitlines( )
					if not self.lines:
						_headers[ self.lines[ 0 ].strip( ) ] += 1
						_footers[ self.lines[ -1 ].strip( ) ] += 1
				
				# Identify candidates for removal
				_head = { line for line, count in _headers.items( ) if
				          count >= min }
				_foot = { line for line, count in _footers.items( ) if
				          count >= min }
				
				# Second pass: clean pages
				for _page in self.pages:
					self.lines = _page.strip( ).splitlines( )
					if not self.lines:
						self.cleaned_pages.append( _page )
						continue
					
					# Remove header
					if self.lines[ 0 ].strip( ) in _head:
						self.lines = self_node.lines[ 1: ]
					
					# Remove footer
					if self.lines and self.lines[ -1 ].strip( ) in _foot:
						self.lines = self_node.lines[ :-1 ]
					
					self.cleaned_pages.append( "\n".join( self.lines ) )
				_retval = self.cleaned_pages
				return _retval
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'remove_headers( self, pages: List[ str ], min: int=3 ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def normalize_text( self, text: str ) -> str:
		"""
	
			Purpose:
			-----------
			Normalizes the path pages path.
			  - Converts pages to lowercase
			  - Removes accented characters (e.g., é -> e)
			  - Removes leading/trailing spaces
			  - Collapses multiple whitespace characters into a single space
	
			Parameters:
			-----------
			- pages : str
				The raw path pages path to be normalized.
	
			Returns:
			--------
			- str
				A normalized, cleaned_lines version of the path path.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.normalized = (unicodedata.normalize( 'NFKD', text )
				                   .encode( 'ascii', 'ignore' ).decode( 'utf-8' ))
				return self.normalized
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'normalize_text( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def lemmatize_tokens( self, tokens: List[ str ] ) -> List[ str ]:
		"""
	
			Purpose:
			-----------
			Performs lemmatization on the path List[ str ] into a path
			of word-tokens.
	
			This function:
			  - Converts pages to lowercase
			  - Tokenizes the lowercased pages into words
			  - Lemmatizes each token using WordNetLemmatizer
			  - Reconstructs the lemmatized tokens into a single path
	
			Parameters:
			-----------
			- pages : str
				The path pages path to be lemmatized.
	
			Returns:
			--------
			- str
				A path with all words lemmatized.
	
		"""
		
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" is required.' )
			else:
				self.tokens = tokens
				pos_tags = nltk.pos_tag( self.tokens )
				self.lemmatized = [ self.lemmatizer.lemmatize( word, get_wordnet_pos( tag ) ) for
				                    word, tag in pos_tags ]
				return self.lemmatized
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'tokenize_words( self, tokens: List[ str  ] ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def tokenize_text( self, text: str ) -> List[ str ]:
		'''

			Purpose:
			---------
			Splits the raw path removes non-words and returns tokens

			Parameters:
			-----------
			- cleaned_line: (str) - clean documents.
	
			Returns:
			- list: Cleaned and normalized documents.
	
		'''
		try:
			if text is None:
				raise Exception( 'The argument "text" was None' )
			else:
				self.tokens = nltk.word_tokenize( text )
				self.words = [ w for w in self.tokens ]
				_retokens = [ re.sub( r'[^\w"-]', '', word ) for word in self.words if
				              word.strip( ) ]
				return _retokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'tokenize_text( self, path: str ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def tiktokenize( self, text: str, encoding: str='cl100k_base' ) -> List[ str ]:
		"""

			Purpose:
			---------
		    Tokenizes text text into subword tokens using OpenAI's tiktoken tokenizer.
		    This function leverages the tiktoken library, which provides byte-pair encoding (BPE)
		    tokenization used in models such as GPT-3.5 and GPT-4. Unlike standard word tokenization,
		    this function splits text into model-specific subword units.
		
		    Parameters
		    ----------
		    - text : str
		        The text string to be tokenized.
		
		    - model : str, optional
		        The tokenizer model to use. Examples include 'cl100k_base' (default),
		        'gpt-3.5-turbo', or 'gpt-4'. Ensure the model is supported by tiktoken.
		
		    Returns
		    -------
		    - List[str]
		        A list of string tokens representing BPE subword units.

        """
		try:
			if text is None:
				raise Exception( 'The argument "text" was None' )
			else:
				self.encoding = tiktoken.get_encoding( encoding )
				_token_ids = self.encoding.encode( text )
				self.tokens = nltk.word_tokenize( text )
				self.words = [ w for w in self.tokens ]
				_retokens = [ re.sub( r'[^\w"-]', '', word ) for word in self.words if
				              word.strip( ) ]
				return _retokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'tiktokenize( self, text: str, encoding: str="cl100k_base" ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def tokenize_words( self, words: List[ str ] ) -> List[ str ]:
		"""
	
			Purpose:
			-----------
			  - Tokenizes the path pages path into individual word tokens.
			  - Converts pages to lowercase
			  - Uses NLTK's word_tokenize to split
			  the pages into words and punctuation tokens
	
			Parameters:
			-----------
			- words : List[ str ]
				A list of strings to be tokenized.
	
			Returns:
			--------
			- A list of token strings (words and punctuation) extracted from the pages.
	
		"""
		try:
			if words is None:
				raise Exception( 'The argument "words" was None' )
			else:
				self.words = words
				for w in self.words:
					_tokens = nltk.word_tokenize( w )
					self.tokens.append( _tokens )
				return self.tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'tokenize_words( self, path: str ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def tokenize_sentences( self, text: str ) -> List[ str ]:
		"""

			Purpose:
			--------
			Tokenize a paragraph or
			document into a list[ str ] of sentence strings.
	
			Parameters:
			-----------
			- text (str): Input pages.
	
			Returns:
			--------
			- list: List of sentence strings.
				
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				self.raw_input = text
				self.tokens = nltk.sent_tokenize( self.raw_input )
				return self.tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'tokenize_sentences( self, text: str ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def chunk_text( self, text: str, size: int=50, return_as_string: bool=True ) -> List[
		List[ str ] ]:
		"""
	
			Purpose:
			-----------
			Tokenizes cleaned_lines pages and breaks it into chunks for downstream vectors.
			  - Converts pages to lowercase
			  - Tokenizes pages using NLTK's word_tokenize
			  - Breaks tokens into chunks of a specified size
			  - Optionally joins tokens into strings (for transformer models)
	
			Parameters:
			-----------
			- pages : str
				The cleaned_lines path pages to be tokenized and chunked.
	
			- chunk_size : int, optional (default=50)
				Number of tokens per chunk_words.
	
			- return_string : bool, optional (default=True)
				If True, returns each chunk_words as a path; otherwise, returns a get_list of
				tokens.
	
			Returns:
			--------
			- a list
				Each chunk_words is either a get_list of tokens or a path.
	
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" is required.' )
			else:
				
				self.tokens = nltk.word_tokenize( text )
				self.chunks = [ self.tokens[ i: i + size ] for i in
				                range( 0, len( self.tokens ), size ) ]
				if return_as_string:
					return [ ' '.join( chunk ) for chunk in self.chunks ]
				else:
					return self.chunks
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'chunk_text( self, text: str, max: int=800 ) -> list[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def chunk_words( self, words: List[ str ], size: int=50, as_string: bool=True ) -> List[
		List[ str ] ]:
		"""
	
			Purpose:
			-----------
			Breaks a list of tokenized strings into a list of lists.
	
			This function:
			  - Groups tokens into chunks of min `chunk_size`
			  - Returns a list of lists of tokens
	
			Parameters:
			-----------
			- words : a list of tokenizd words
	
			- chunk_size : int, optional (default=50)
				Number of tokens per chunk_words.
	
			Returns:
			--------
			- List[ List[ str ] ]
				A list of a list of token chunks. Each chunk is a list of tokens.
	
		"""
		try:
			if words is None:
				raise Exception( 'The argument "words" is required.' )
			else:
				self.tokens = [ token for sublist in words for token in sublist ]
				self.chunks = [ self.tokens[ i: i + size ]
				                for i in range( 0, len( self.tokens ), size ) ]
				if as_string:
					return [ ' '.join( chunk ) for chunk in self.chunks ]
				else:
					return self.chunks
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = (
				'chunk_words( self, tokens: list[ str ], max: int=800, over: int=50 ) -> list[ '
				'str ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def split_sentences( self, text: str ) -> List[ str ]:
		"""

			Purpose:
			________
			Splits the text text string into a list of
			individual sentences using NLTK's Punkt sentence tokenizer.
			This function is useful for preparing text for further linguistic processing,
			such as tokenization, parsing, or named entity recognition.

			Parameters
			----------
			- text : str
				The raw text string to be segmented into sentences.

			Returns
			-------
			- List[str]
				A list of sentence strings, each corresponding to a single sentence detected
				in the text text.

		"""
		try:
			return nltk.sent_tokenize( text )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'split_sentences( self, text: str ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def split_pages( self, path: str, delimit: str='\f' ) -> List[ str ]:
		"""

			Purpose:
			--------
			Reads path from a file, splits it into tokens,
			and groups them into path.
	
			Parameters:
			-----------
			- path (str): Path to the path file.
			- delimiter (str): Page separator path (default is '\f' for form feed).
	
			Returns:
			---------
			- list[ str ]  where each element is the path.
	
		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" is required' )
			else:
				self.file_path = path
				with open( self.file_path, 'r', encoding='utf-8', errors='ignore' ) as _file:
					_content = _file.read( )
				self.raw_pages = _content.split( delimit )
				for _page in self.raw_pages:
					self.lines = _page.strip( ).splitlines( )
					self.cleaned_text = '\n'.join( [ line.strip( )
					                                 for line in self.lines if line.strip( ) ] )
					self.cleaned_pages.append( self.cleaned_text )
				return self.cleaned_pages
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'split_pages( self, path: str, delimit: str="\f" ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def split_paragraphs( self, path: str ) -> List[ str ]:
		"""

			Purpose:
			---------
			Reads  a file and splits it into paragraphs. A paragraph is defined as a block
			of path separated by one or more empty tokens.
	
			Parameters:
			-----------
			- path (str): Path to the path file.
	
			Returns:
			---------
			- list of str: List of paragraph strings.
				
		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" is required.' )
			else:
				self.file_path = path
				with open( self.file_path, 'r', encoding='utf-8' ) as _file:
					self.raw_input = _file.read( )
					self.paragraphs = [ para.strip( ) for para in self.raw_input.split( '\n\n' ) if
					                    para.strip( ) ]
					
					return self.paragraphs
		except UnicodeDecodeError:
			with open( self.file_path, 'r', encoding='latin1' ) as _file:
				self.raw_input = _file.read( )
				self.paragraphs = [ para.strip( ) for para in self.raw_input.split( '\n\n' ) if
				                    para.strip( ) ]
				return self.paragraphs
	
	
	def compute_frequency_distribution( self, lines: List[ str ], process: bool=True ) -> FreqDist:
		"""

			Purpose:
			--------
			Creates a word frequency freq_dist from a list of documents.
	
			Parameters:
			-----------
			- documents (list): List of raw or preprocessed path documents.
			- process (bool): Applies normalization, tokenization, stopword removal, and lemmatization.
	
			Returns:
			- dict: Dictionary of words and their corresponding frequencies.
				
		"""
		try:
			if lines is None:
				raise Exception( 'The argument "lines" is required.' )
			else:
				self.lines = lines
				for _line in self.lines:
					if process:
						self.normalized = self.normalize_text( _line )
						self.words = self.tokenize_words( self.normalized )
						self.tokens = self.lemmatize_tokens( self.words )
					else:
						self.words = self.tokenize_words( _line )
						self.tokens.append( self.words )
				self.frequency_distribution = dict( Counter( self.tokens ) )
				return self.frequency_distribution
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = ('compute_frequency_distribution( self, documents: list, process: '
			               'bool=True) -> FreqDist')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def compute_conditional_distribution( self, lines: List[ str ], condition=None,
	                                      process: bool=True ) -> ConditionalFreqDist:
		"""

			Purpose:
			--------
			Computes a Conditional Frequency Distribution (CFD)
			 over a collection of documents.
	
			Parameters:
			-----------
			- documents (list):
				A list of path sections (pages, paragraphs, etc.).
	
			- condition (function):
				A function to determine the condition/grouping. If None, uses document index.
	
			- process (bool):
				If True, applies normalization, tokenization,
				stopword removal, and lemmatization.
	
			Returns:
			- ConditionalFreqDist:
				An NLTK ConditionalFreqDist object mapping conditions to word frequencies.
	
		"""
		try:
			if lines is None:
				raise Exception( 'The argument "lines" is required.' )
			else:
				self.lines = lines
				self.conditional_distribution = ConditionalFreqDist( )
				for idx, _line in enumerate( self.lines ):
					condition = condition( _line ) if condition else f'Doc_{idx}'
					if process:
						self.normalized = self.normalize_text( _line )
						self.words = self.tokenize_words( self.normalized )
						self.tokens = self.lemmatize_tokens( self.words )
					else:
						self.tokens = self.tokenize_words( _line )
					
					for _token in self.tokens:
						self.conditional_distribution[ condition ][ _token ] += 1
				
				return self.conditional_distribution
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = ('compute_conditional_distribution( self, tokens: List[ str ], '
			               'condition=None, process: bool=True ) -> ConditionalFreqDist')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_vocabulary( self, freq_dist: Dict, min: int=1 ) -> List[ str ]:
		"""

			Purpose:
			---------
			Builds a vocabulary list from a frequency
			distribution by applying a minimum frequency threshold.
	
			Parameters:
			-----------
			- freq_dist (dict):
				A dictionary mapping words to their frequencies.

			- min (int): Minimum num
				of occurrences required for a word to be included.
	
			Returns:
			--------
			- list: Sorted list of unique vocabulary words.
				
		"""
		try:
			if freq_dist is None:
				raise Exception( 'The argument "freq_dist" is required.' )
			else:
				self.frequency_distribution = freq_dist
				self.words = [ word for word, freq in freq_dist.items( ) if freq >= min ]
				self.vocabulary = sorted( self.words )
				return self.vocabulary
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'create_vocabulary( self, freq_dist: dict, min: int=1 ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_wordbag( self, tokens: List[ str ] ) -> dict:
		"""
			
			Purpose:
			--------
			Construct a Bag-of-Words (BoW) frequency dictionary from a list of strings.
	
			Parameters:
			-----------
			- tokens (list): List of tokens from a document.
	
			Returns:
			--------
			- dict: Word frequency dictionary.
				
		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" is required.' )
			else:
				self.tokens = tokens
				return dict( Counter( self.tokens ) )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'create_wordbag( self, tokens: List[ str ] ) -> dict'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_word2vec( self, tokens: List[ str ], size=100, window=5, min=1 ) -> Word2Vec:
		"""

			Purpose:
			--------
			Train a Word2Vec embedding small_model from tokenized sentences.
	
			Parameters:
			--------
			- sentences (get_list of get_list of str): List of tokenized sentences.
			- vector_size (int): Dimensionality of word vec.
			- window (int): Max distance between current and predicted word.
			- min_count (int): Minimum frequency for inclusion in vocabulary.
	
			Returns:
			-------
			- Word2Vec: Trained Gensim Word2Vec small_model.

		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" is required.' )
			else:
				self.tokens = tokens
				return Word2Vec( sentences=self.tokens, vector_size=size,
					window=window, min_count=min )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = ('create_word2vec( self, tokens: list, '
			               'size=100, window=5, min=1 ) -> Word2Vec')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_tfidf( self, lines: List[ str ], max: int=1000, prep: bool=True ) -> Tuple:
		"""
		
			Purpose:
			________
			Compute TF-IDF matrix with
			optional full preprocessing pipeline.
	
			Parameters:
			--------
			- lines (list): List of raw or preprocessed pages documents.
			- max (int): Max num of terms to include (vocabulary size).
			- prep (bool): If True, normalize, tokenize_text, clean, and lemmatize path.
	
			Returns:
			--------
			- tuple:
			- tfidf_matrix (scipy.sparse.csr_matrix): TF-IDF feature matrix.
			- features (list): Vocabulary terms.
			- vectorizer (TfidfVectorizer): Fitted vectorizer instance.

		"""
		try:
			if lines is None:
				raise Exception( 'The argument "tokens" is required.' )
			elif prep:
				self.lines = lines
				for _doc in self.lines:
					self.normalized = self.normalize( _doc )
					self.tokens = self.tokenize_words( self.normalized )
					self.words = [ self.lemmatize( token ) for token in self.tokens ]
					self.cleaned_text = ' '.join( self.words )
					self.cleaned_lines.append( cleaned_text )
				self.vectorizer = TfidfVectorizer( max_features=max, stop_words='english' )
				_matrix = self.vectorizer.fit_transform( self.cleaned_lines )
				return tuple( _matrix, self.vectorizer.get_feature_names_out( ).tolist( ), self.vectorizer )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = ('create_tfidf( self, tokens: list, max: int=1000, prep: bool=True ) '
			               '-> Tuple' )
			error = ErrorDialog( exception )
			error.show( )
	
	
	def clean_files( self, src: str, dest: str ) -> None:
		"""
		
			Purpose:
			________
			Cleans text files given a source directory (src) and destination directory (dest)
	
			Parameters:
			----------
			- src (str): Source directory
			- dest (str): Destination directory
	
			Returns:
			--------
			- None
	
		"""
		try:
			if src is None:
				raise Exception( 'The argument "src" is required.' )
			elif dest is None:
				raise Exception( 'The argument "dest" is required.' )
			else:
				source = src
				destination = dest
				files = os.listdir( source )
				for f in files:
					processed = [ ]
					filename = os.path.basename( f )
					source_path = source + '\\' + filename
					text = open( source_path, 'r', encoding='utf-8', errors='ignore' ).read( )
					punc = self.remove_punctuation( text )
					dirty = self.split_sentences( punc )
					for d in dirty:
						if d != " ":
							lower = d.lower( )
							normal = self.normalize_text( lower )
							spec = self.remove_special( normal )
							slim = self.collapse_space( spec )
							processed.append( slim )
					
					dest_path = destination + '\\' + filename
					clean = open( dest_path, 'wt', encoding='utf-8', errors='ignore' )
					for p in processed:
						clean.write( p )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'clean_files( self, src: str, dest: str )'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def convert_jsonl( self, source: str, desination: str ) -> None:
		"""
		
			Purpose:
			________
			Coverts text files to JSONL format given a source directory (Source)
			 and destination directory (destination)
	
			Parameters:
			--------
			- source (str): Source directory
			- destination (str): Destination directory
	
			Returns:
			--------
			- None
	
		"""
		try:
			if source is None:
				raise Exception( 'The argument "sourc" is required.' )
			elif desination is None:
				raise Exception( 'The argument "desination" is required.' )
			else:
				_source = source
				_destination = desination
				self.files = os.listdir( _source )
				_processed = [ ]
				for _f in self.files:
					_count = 0
					_basename = os.path.basename( _f )
					_sourcepath = _source + f'\\{_basename}'
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_stops = self.remove_stopwords( _text )
					_tokens = self.tokenize_text( _stops )
					_chunks = self.chunk_text( _text )
					_filename = _basename.rstrip( '.txt' )
					_destinationpath = _destination + f'\\{_filename}.jsonl'
					_clean = open( _destinationpath, 'wt', encoding='utf-8', errors='ignore' )
					for _i in range( len( _chunks ) ):
						_list = _chunks[ _i ]
						_part = ''.join( _list )
						_row = '{' + f'\"Line-{_i}\":\"{_part}\"' + '}' + '\r'
						_processed.append( _row )
					
					for _t in _processed:
						_clean.write( _t )
					
					_clean.flush( )
					_clean.close( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Text'
			exception.method = 'convert_jsonl( self, source: str, desination: str )'
			error = ErrorDialog( exception )
			error.show( )


class Word( BaseModel ):
    """

        Purpose:
        --------
        A class to extract, clean, and analyze text from Microsoft Word documents.

        Methods:
        --------
        split_sentences( self ) -> None
        clean_sentences( self ) -> None
        create_vocabulary( self ) -> None
        compute_frequency_distribution( self ) -> None
        summarize( self ) -> None:

    """
    filepath: str
    raw_text: str
    paragraphs: Optionl[ List[ str ] ]
    sentences: Optional[ List[ str ] ]
    cleaned_sentences: Optional[ List[ str ] ]
    vocabulary: Optional[ List[ str ] ]
    freq_dist: Optional[ Dict[ str, float ] ]


    class Config:
	    arbitrary_types_allowed = True
	    extra = 'ignore'
	    allow_mutation = True


    def __init__( self, filepath: str ) -> None:
        """

            Purpose:
            --------
            Initializes the WordTextProcessor with the path to the .docx file.

			Parameters:
			----------
            :param filepath: Path to the Microsoft Word document (.docx)

        """
        super( ).__init__( )
        self.filepath = filepath
        self.raw_text = ''
        self.paragraphs = [ ]
        self.sentences = [ ]
        self.cleaned_sentences = [ ]
        self.vocabulary = [ ]
        self.freq_dist = { }


    def __dir__( self ) -> List[ str ]:
        '''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
        return [ 'extract_text', 'split_sentences', 'clean_sentences',
                 'create_vocabulary', 'compute_frequency_distribution',
                 'summarize', 'filepath', 'raw_text', 'paragraphs',
                 'sentences', 'cleaned_sentences', 'vocabulary', 'freq_dist' ]


    def extract_text( self ) -> None:
        """

            Purpose:
            --------
            Extracts raw text and paragraphs from the .docx file.

        """
        document = Docx( self.filepath )
        self.paragraphs = [ para.text.strip( ) for para in document.paragraphs if para.text.strip( ) ]
        self.raw_text = '\n'.join( self.paragraphs )


    def split_sentences( self ) -> None:
        """

            Purpose:
            --------
            Splits the raw text into sentences.

        """
        try:
	        self.sentences = sent_tokenize( self.raw_text )
        except Exception as e:
	        exception = Error( e )
	        exception.module = 'Tiggr'
	        exception.cause = 'Word'
	        exception.method = ''
	        error = ErrorDialog( exception )
	        error.show( )


    def clean_sentences( self ) -> None:
        """

            Purpose:
            -------
            Cleans each sentence: removes extra whitespace, punctuation, and lowers the text.

        """
        try:
	        self.cleaned_sentences = [ ]
	        for sentence in self.sentences:
		        sentence = re.sub( r"[\r\n\t]+", " ", sentence )
		        sentence = re.sub( r"[^a-zA-Z0-9\s']", "", sentence )
		        sentence = re.sub( r"\s{2,}", " ", sentence ).strip( ).lower( )
		        self.cleaned_sentences.append( sentence )
        except Exception as e:
	        exception = Error( e )
	        exception.module = 'Tiggr'
	        exception.cause = 'Word'
	        exception.method = ''
	        error = ErrorDialog( exception )
	        error.show( )


    def create_vocabulary( self ) -> None:
        """

            Purpose:
            --------
            Computes vocabulary terms from cleaned sentences.

        """
        try:
	        all_words = [ ]
	        stop_words = set( stopwords.words( 'english' ) )
	        for sentence in self.cleaned_sentences:
		        tokens = word_tokenize( sentence )
		        tokens = [ token for token in tokens if
		                   token.isalpha( ) and token not in stop_words ]
		        all_words.extend( tokens )
	        self.vocabulary = sorted( set( all_words ) )
        except Exception as e:
	        exception = Error( e )
	        exception.module = 'Tiggr'
	        exception.cause = 'Word'
	        exception.method = ''
	        error = ErrorDialog( exception )
	        error.show( )


    def compute_frequency_distribution( self ) -> None:
        """

            Purpose:
            -------
            Computes frequency distribution of the vocabulary.

        """
        try:
	        words = [ ]
	        for sentence in self.cleaned_sentences:
		        tokens = word_tokenize( sentence )
		        tokens = [ token for token in tokens if token.isalpha( ) ]
		        words.extend( tokens )
	        self.freq_dist = dict( Counter( words ) )
        except Exception as e:
	        exception = Error( e )
	        exception.module = 'Tiggr'
	        exception.cause = 'Word'
	        exception.method = 'correct_errors( self, text: str ) -> str'
	        error = ErrorDialog( exception )
	        error.show( )


    def summarize( self ) -> None:
        """

            Purpose:
            -------
            Prints a summary of extracted and processed text.

        """
        print( f'Document: {self.filepath}' )
        print( f'Paragraphs: {len( self.paragraphs )}' )
        print( f'Sentences: {len( self.sentences )}' )
        print( f'Vocabulary Size: {len( self.vocabulary )}' )
        print( f'Top 10 Frequent Words: { Counter( self.freq_dist ).most_common(10)}' )

	
class PDF( BaseModel ):
	"""

		Purpose:
		--------
		A utility class for extracting clean pages from PDF files into a list of strings.
		Handles nuances such as layout artifacts, page separation, optional filtering,
		and includes df detection capabilities.


	    Methods:
	    --------
	    extract_lines( self, path, max: int=None) -> List[ str ]
	    extract_text( self, path, max: int=None) -> str
	    export_csv( self, tables: List[ pd.DataFrame ], filename: str=None ) -> None
	    export_text( self, lines: List[ str ], path: str=None ) -> None
	    export_excel( self, tables: List[ pd.DataFrame ], path: str=None ) -> None

	"""


	def __init__( self, headers: bool=False, min: int=10, tables: bool=True ) -> None:
		"""

			Purpose:
			-----------
			Initialize the PDF pages extractor with configurable settings.

			Parameters:
			-----------
			- headers (bool): If True, attempts to strip recurring headers/footers.
			- min (int): Minimum num of characters for a line to be included.
			- tables (bool): If True, extract pages from detected tables using block
			grouping.

		"""
		super( ).__init__( )
		self.strip_headers = headers
		self.minimum_length = min
		self.extract_tables = tables
		self.pages = [ ]
		self.lines = [ ]
		self.clean_lines = [ ]
		self.extracted_lines = [ ]
		self.extracted_tables = [ ]
		self.extracted_pages = [ ]
		self.file_path = None
		self.page = None


	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
		return [ 'strip_headers', 'minimum_length', 'extract_tables',
		         'path', 'page', 'pages', 'tokens', 'clean_lines', 'extracted_lines',
		         'extracted_tables', 'extracted_pages', 'extract_lines',
		         'extract_text', 'extract_tables', 'export_csv',
		         'export_text', 'export_excel' ]


	def extract_lines( self, path: str, max: Optional[ int ]=None ) -> List[ str ]:
		"""

			Purpose:
			--------
			Extract tokens of pages from a PDF,
			optionally limiting to the first N pages.

			Parameters:
			----------
			- path (str): Path to the PDF file
			- max (Optional[int]): Max num of pages to process (None for all pages)

			Returns:
			--------
			- List[str]: Cleaned list of non-empty tokens

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be specified' )
			else:
				self.file_path = path
				with fitz.open( self.file_path ) as doc:
					for i, page in enumerate( doc ):
						if max is not None and i >= max:
							break
						if self.extract_tables:
							self.extracted_lines = self._extract_tables( page )
						else:
							_text = page.get_text( 'pages' )
							self.lines = _text.splitlines( )
						self.clean_lines = self._filter_lines( self.lines )
						self.extracted_lines.extend( self.clean_lines )
				return self.extracted_lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = ('extract_lines( self, path: str, max: Optional[ int ]=None ) -> '
			               'List[ str ]')
			error = ErrorDialog( exception )
			error.show( )

	
	def _extract_tables( self, page: Page ) -> List[ str ]:
		"""

			Purpose:
			--------
			Attempt to extract structured blocks
			such as tables using spatial grouping.

			Parameters:
			----------
			- page: PyMuPDF page object

			Returns:
			--------
			- List[str]: Grouped blocks including potential tables

		"""
		try:
			if page is None:
				raise Exception( 'The argument "page" cannot be None' )
			else:
				_blocks = page.get_text( 'blocks' )
				_sorted = sorted( _blocks, key=lambda b: (round( b[ 1 ], 1 ), round( b[ 0 ], 1 ) ) )
				self.lines = [ b[ 4 ].strip( ) for b in _sorted if b[ 4 ].strip( ) ]
				return self.lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = '_extract_tables( self, page ) -> List[ str ]:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def _filter_lines( self, lines: List[ str ] ) -> List[ str ]:
		"""

			Purpose:
			-----------
			Filter and clean tokens from a page of pages.

			Parameters:
			- tokens (List[str]): Raw tokens of pages

			Returns:
			--------
			- List[str]: Filtered, non-trivial tokens

		"""
		try:
			if line is None:
				raise Exception( 'The argument "line" is None' )
			else:
				self.lines = lines
				for line in self.lines:
					_line = line.strip( )
					if len( _line ) < self.minimum_length:
						continue
					if self.strip_headers and self._has_repeating_header( _line ):
						continue
					self.clean_lines.append( _line )
				return self.clean_lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = '_filter_lines( self, tokens: List[ str ] ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def _has_repeating_header( self, line: str ) -> bool:
		"""

			Purpose:
			--------
			Heuristic to detect common headers/footers (basic implementation).

			Parameters:
			- line (str): A line of pages

			Returns:
			--------
			- bool: True if line is likely a header or footer

		"""
		try:
			if line is None:
				raise Exception( 'The argument "line" is None' )
			else:
				_keywords = [ 'page', 'public law', 'u.s. government', 'united states' ]
				return any( kw in line.lower( ) for kw in _keywords )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = '_has_repeating_header( self, line: str ) -> bool'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def extract_text( self, path: str, max: Optional[ int ]=None ) -> str:
		"""

			Purpose:
			---------
			Extract the entire pages from a PDF into one continuous path.

			Parameters:
			-----------
			- path (str): Path to the PDF file
			- max (Optional[int]): Maximum num of pages to process

			Returns:
			--------
			- str: Full concatenated pages

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be specified' )
			else:
				if max is not None and max > 0:
					self.file_path = path
					self.lines = self.extract_lines( self.file_path, max=max )
					return '\n'.join( self.lines )
				elif max is None or max <= 0:
					self.file_path = path
					self.lines = self.extract_lines( self.file_path )
					return '\n'.join( self.lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = 'extract_text( self, path: str, max: Optional[ int ]=None ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def extract_tables( self, path: str, max: Optional[ int ]=None ) -> List[ pd.DataFrame ]:
		"""

			Purpose:
			-----------
			Extract tables from the PDF and return them as a list of DataFrames.

			Parameters:
			- path (str): Path to the PDF file
			- max (Optional[int]): Maximum num of pages to process

			Returns:
			--------
			- List[pd.DataFrame]: List of DataFrames representing detected tables

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be specified' )
			elif max is None:
				raise Exception( 'The argument "max" must be specified' )
			else:
				self.file_path = path
				with fitz.open( self.file_path ) as _doc:
					for i, page in enumerate( _doc ):
						if max is not None and i >= max:
							break
						_blocks = page.find_tables( )
						for _tables in _blocks.tables:
							_dataframe = pd.DataFrame( _tables.extract( ) )
							self.tables.append( _dataframe )
				return self.tables
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = (
				'extract_tables( self, path: str, max: Optional[ int ] = None ) -> List[ '
				'pd.DataFrame ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def export_csv( self, tables: List[ pd.DataFrame ], filename: str ) -> None:
		"""

			Purpose:
			-----------
			Export a list of DataFrames (tables) to individual CSV files.

			Parameters:
			- tables (List[pd.DataFrame]): List of tables to export
			- filename (str): Prefix for output filenames (e.g., 'output_table')

		"""
		try:
			if tables is None:
				raise Exception( 'The argument "tables" must not be None' )
			elif filename is None:
				raise Exception( 'The argument "filename" must not be None' )
			else:
				self.tables = tables
				for i, df in enumerate( self.tables ):
					df.to_csv( f'{filename}_{i + 1}.csv', index=False )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = 'export_csv( self, tables: List[ pd.DataFrame ], filename: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def export_text( self, lines: List[ str ], path: str ) -> None:
		"""

			Export extracted tokens of
			pages to a plain pages file.

			Parameters:
			-----------
			- tokens (List[str]): List of pages tokens
			- path (str): Path to output pages file

		"""
		try:
			if lines is None:
				raise Exception( 'The argument "tokens" must be provided.' )
			elif path is None:
				raise Exception( 'The argument "path" must be provided.' )
			else:
				self.file_path = path
				self.lines = lines
				with open( self.file_path, 'w', encoding='utf-8' ) as f:
					for line in self.lines:
						f.write( line + '\n' )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = 'export_text( self, tokens: List[ str ], path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def export_excel( self, tables: List[ pd.DataFrame ], path: str ) -> None:
		"""

			Export all extracted tables into a single
			Excel workbook with one sheet per df.

			Parameters:
			-----------
			- tables (List[pd.DataFrame]): List of tables to export
			- path (str): Path to the output Excel file

		"""
		try:
			if tables is None:
				raise Exception( 'The argument "tables" must not be None' )
			elif path is None:
				raise Exception( 'The argument "path" must not be None' )
			else:
				self.tables = tables
				self.file_path = path
				with pd.ExcelWriter( self.file_path, engine='xlsxwriter' ) as _writer:
					for i, df in enumerate( self.tables ):
						_sheet = f'Table_{i + 1}'
						df.to_excel( writer, sheet_name=_sheet, index=False )
					_writer.save( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'PDF'
			exception.method = 'export_excel( self, tables: List[ pd.DataFrame ], path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )


class Token( BaseModel ):
	'''
	
	
		Purpose:
		________
		
		Wrapper for Hugging Face tokenizers using the `transformers` library.
	    Includes sentence-level segmentation, tokenization, encoding, and decoding.
	    
	    Methods:
	    _______
	    
	    encode( self, path: str ) -> List[str]
	    batch_encode( self, path: str ) -> List[str]
	    decode( self, ids: List[ str ], skip: bool=True ) -> List[str]
	    convert_tokens( self, tokens: List[str] ) -> List[str]
	    convert_ids( self, ids: List[str] ) -> List[str]
	    get_vocab( self ) -> List[str]
	    save_tokenizer( self, path: str ) -> None
	    load_tokenizer( self, path: str ) -> None
	
	'''
	encoding: Optional[ Encoding ]
	model_name: Optional[ str ]
	raw_input: Optional[ str ]
	tokenizer: Optional[ object ]

	class Config:
		arbitrary_types_allowed = True
		extra = 'ignore'
		allow_mutation = True


	def __init__( self ):
		'''
		
		
			Purpose:
	        Initializes the tokenizer wrapper using a pre-trained small_model from Hugging Face.
	
			Parameters:
			--------
	            model_name (str): The name of the pre-trained small_model (e.g.,
	            "bert-base-uncased").
        '''
		super( ).__init__( )
		self.model_name = 'google-bert/bert-base-uncased'
		self.tokenizer = AutoTokenizer.from_pretrained( self.model_name, trust_remote_code=True )
		self.raw_input = None
		self.encoding = None


	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
		return [ 'raw_input', 'encoding', 'tokenizer', 'model_name',
		         'tiktoken_count', 'get_vocab', 'load_tokenizer',
		         'save_tokenizer', 'encode', 'batch_encode', 'convert_tokens',
		         'convert_ids', 'decode' ]


	def tiktoken_count( self, text: str, encoding: str='cl100k_base' ) -> int:
		"""
		
			Purpose:
			_______
			Method counts the number of tokens in a string of path.
			
			Parameters:
			__________
			path - a string given as text
			encoding - the encoding scaler
			
			Returns:
			_______
			Returns the number of tokens in a path path.
			
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must not be None' )
			elif encoding is None:
				raise Exception( 'The argument "encoding" must not be None' )
			else:
				self.encoding = tiktoken.get_encoding( encoding )
				num_tokens = len( encoding.encode( text ) )
				return num_tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'tiktoken_count( self, path: str, encoding: str ) -> int:'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def get_vocab( self ) -> Dict[ str, int ]:
		"""
			
			Retrieves the
			tokenizer's vocabulary.
	
			Returns:
			--------
				Dict[str, int]: Mapping of token path to token ID.
			
		"""
		try:
			return self.tokenizer.get_vocab( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'get_vocab( self ) -> Dict[ str, int ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def load_tokenizer( self, path: str ) -> None:
		"""

			Purpose:
			-------
			Loads a tokenizer from
			 a specified directory path.
	
			Parameters:
			--------
				path (str): Path to the tokenizer config and vocab files.
			
		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be provided.' )
			else:
				self.tokenizer = AutoTokenizer.from_pretrained( path )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'load_tokenizer( self, path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def save_tokenizer( self, path: str ) -> None:
		"""

			Purpose:
			--------
			Saves the tokenizer
			to a directory.
	
			Parameters:
			--------
				path (str): Target path to save tokenizer config and vocab.
			
		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be provided.' )
			else:
				self.tokenizer.save_pretrained( path )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'save_tokenizer( self, path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def encode( self, text: str, max: int=512, trunc: bool=True,
	            padd: Union[ bool, str ]=False, tensors: str=None ) -> Dict[
		str, Union[ List[ int ], any ] ]:
		"""

			Purpose:
			---------
			Encodes a single path of path
			into small_model-ready path IDs and attention masks.
	
			Parameters:
			----------
				text (str): Input path.
				max (int): Max length of token sequence.
				trunc (bool): If True, trunc sequences over max.
				padd (bool | str): If True or 'max', pad to max length.
				tensors (str): One of 'pt', 'tf', or 'np'.
	
			Returns:
			--------
				Dict[str, any]: Dictionary with input_ids, attention_mask, etc.
			
		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must be provided.' )
			else:
				return self.tokenizer( text, truncation=trunc, padding=padd,
					max_length=max, return_tensors=tensors )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'encode( self, text: str ) -> Dict[ str, Union[ List[ int ], any ] ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def batch_encode( self, texts: List[ str ], max: int=512, trunc: bool=True,
	                  pad: Union[ bool, str ]='max', tensors: str=None ) -> Dict[ str, any ]:
		"""
			
			Encodes a list of
			path inputs as a batch.
	
			Parameters:
			-----------
				texts (List[str]): A list of path samples.
				max (int): Max length for truncate.
				trunc (bool): Whether to truncate.
				pad (bool | str): Padding mode.
				tensors (str): Output tensor scaler.
	
			Returns:
			--------
				Dict[str, any]: Tokenized batch with path IDs, masks, etc.
			
		"""
		try:
			if texts is None:
				raise Exception( 'The argument "texts" must be provided.' )
			else:
				return self.tokenizer( texts, truncation=trunc, adding=pad,
					max_length=max, return_tensors=tensors )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = '''batch_encode( self, texts: List[ str ], max: int=512,
			trunc: bool=True,
	                  pad: Union[ bool, str ]='max', tensors: str=None ) -> Dict[ str, any ]'''
			error = ErrorDialog( exception )
			error.show( )
	
	
	def decode( self, ids: List[ int ], skip: bool=True ) -> str:
		"""

			Purpose:
			--------
			Converts a list of token IDs back to a path.
	
			Parameters:
			-----------
				ids (List[int]): Encoded token IDs.
				skip (bool): Exclude special tokens from output.
	
			Returns:
			--------
				str: Human-readable decoded path.
				
		"""
		try:
			if ids is None:
				raise Exception( 'The argument "ids" must be provided.' )
			else:
				return self.tokenizer.decode( ids, skip_special_tokens=skip )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'decode( self, ids: List[ int ], skip: bool=True ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def convert_tokens( self, tokens: List[ str ] ) -> List[ int ]:
		"""

			Purpose:
			--------
			Converts tokens into their corresponding vocabulary IDs.
	
			Parameters:
			-----------
				tokens (List[str]): List of subword tokens.
	
			Returns:
			--------
				List[int]: Token IDs.
			
		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" must be provided.' )
			else:
				return self.tokenizer.convert_tokens_to_ids( tokens )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'convert_tokens( self, tokens: List[ str ] ) -> List[ int ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def convert_ids( self, ids: List[ int ] ) -> List[ str ]:
		"""

			Purpose:
			--------
			Converts token IDs
			back to subword tokens.
	
			Parameters:
			-----------
				ids (List[int]): List of token IDs.
	
			Returns:
			--------
				List[str]: List of token strings.
			
		"""
		try:
			if ids is None:
				raise Exception( 'The argument "ids" must be provided.' )
			else:
				return self.tokenizer.convert_ids_to_tokens( ids )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Token'
			exception.method = 'convert_ids( self, ids: List[ int ] ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )


class VectorStore( BaseModel ):
	"""

		Purpose:
		---------
		A class for generating OpenAI vectors, performing normalization, computing similarity,
		and interacting with OpenAI VectorStore Stores via the OpenAI API. Includes local
		export/import, vector diagnostics, and bulk querying functionality.

	"""
	small_model: str
	large_model: str
	ada_model: str
	tables: Optional[ List[ pd.DataFrame ] ]
	vectors: Optional[ List[ List[ float ] ] ]
	array: Optional[ List[ float ] ]
	batches: List[ List[ str ] ]
	vector_stores: Dict[ str, str ]
	store_ids: List[ str ]
	file_ids: List[ str ]
	tokens: List[ str ]
	files: List[ str ]
	raw_text: Optional[ str ]
	file_path: Optional[ str ]
	file_name: Optional[ str ]
	file_ids: Optional[ str ]
	directory: Optional[ str ]
	id: Optional[ str ]
	response: Optional[ object ]
	dataframe: List[ pd.DataFram ]
	client: OpenAI
	cache: Optional[ Dict ]
	stats: Optional[ Dict ]
	results: Optional[ Dict ]
	
	def __init__( self ):
		"""

			Purpose:
			---------
			Initialize the VectorStore object with
			OpenAI API credentials and embedding small_model.

			Parameters:
			-----------
			- api_key (Optional[str]): OpenAI API key (uses global config if None)
			- small_model (str): OpenAI embedding small_model to use

		"""
		super( ).__init__( )
		self.small_model = 'text-embedding-3-small'
		self.large_model = 'text-embedding-3-large'
		self.ada_model = 'text-embedding-ada-002'
		self.client = OpenAI( )
		self.cache = { }
		self.results = { }
		self.stats = { }
		self.data = { }
		self.vector_stores = { }
		self.store_ids = [ ]
		self.file_ids = [ ]
		self.files = [ ]
		self.tokens = [ ]
		self.array = [ ]
		self.vectors = [ [ ] ]
		self.batches = [ [ ] ]
		self.tables = [ ]
		self.raw_text = None
		self.file_path = None
		self.file_name = None
		self.file_ids = None
		self.directory = None
		self.id = None
		self.response = None
		self.dataframe = None
	
	
	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			self

			Returns:
			--------
			List[ str ] | None

		'''
		return [ 'small_model', 'large_model', 'ada_model',
		         'id', 'files', 'tokens', 'array', 'store_ids',
		         'client', 'cache', 'results', 'directory', 'stats',
		         'response', 'vectorstores','file_ids',
		         'df', 'batches', 'tables', 'vectors',
		         'create_small_embedding', 'dataframe',
		         'most_similar', 'bulk_similar', 'similarity_heatmap',
		         'export_jsonl', 'import_jsonl', 'create_vector_store',
		         'list_vector_stores', 'upload_vector_store',
		         'query_vector_store', 'delete_vector_store',
		         'upload_document', 'upload_documents' ]
	
	
	def create( self, tokens: List[ str ], batch: int=10, max: int=3,
	            time: float=2.0 ) -> pd.DataFrame:
		"""

			Purpose:
			---------
			Generate and normalize vectors for a list of path tokens.

			Parameters:
			-----------
			- tokens (List[str]): List of path pages strings
			- batch (int): Number of tokens per API request batch
			- max (int): Number of retries on API failure
			- time (float): Seconds to wait between retries

			Returns:
			--------
			- pd.DataFrame: DataFrame containin
			g original pages, raw vectors,
			and normalized vectors

		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" cannot be None' )
			else:
				self.tokens = tokens
				self.batches = self._batch_chunks( self.tokens, batch )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				for index, batch in enumerate( self.batches ):
					for attempt in range( max ):
						try:
							self.response = self.client.embeddings.create( input=batch,
								model=self.small_model )
							_vectors = [ record.embedding for record in self.response.data ]
							self.vectors.extend( _vectors )
							break
						except Exception as e:
							print( f'[Batch {index + 1}] Retry {attempt + 1}/{max}: {e}' )
							time.sleep( time )
					else:
						raise RuntimeError( f'Failed after {max} attempts on batch {index + 1}' )
				
				_embeddings = np.array( self.array )
				_normed = self._normalize( _embeddings )
				self.data = \
				{
					'pages': tokens,
					'embedding': list( _embeddings ),
					'normed_embedding': list( _normed )
				}
				
				return pd.DataFrame( self.data )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = (
				'create_small_embedding( self, tokens: List[ str ], batch: int=10, max: int=3, '
				'time: float=2.0 ) -> pd.DataFrame')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def _batch_chunks( self, texts: List[ str ], size: int ) -> List[ List[ str ] ]:
		"""


			Purpose:
			---------
			Split a list of tokens into batches of specified size.

			Parameters:
			-----------
			- tokens (List[str]): Full list of path strings
			- size (int): Desired batch size

			Returns:
			--------
			- List of pages batches

		"""
		try:
			if texts is None:
				raise Exception( 'The argument "tokens" cannot be None' )
			elif size is None:
				raise Exception( 'The argument "size" cannot be None' )
			else:
				return [ texts[ i:i + size ] for i in range( 0, len( texts ), size ) ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = (' _batch_chunks( self, tokens: List[ str ], size: int ) -> [ List[ str '
			               '] ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def get_purpose_options( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Returns a list of path representing the purpose of the file

			Parameters:
			-----------
			self

			Returns:
			_________
			List[ str ] | None

		'''
		return [ 'assistants', 'assistants_output', 'batch',
		         'batch_output', 'fine-tune', 'fine-tune-results',
		         'vision' ]
	
	
	def get_type_options( self ) -> Dict[ str, str ] | None:
		'''

			Purpose:
			---------
			Returns a dictionary of file formats and types

			Parameters:
			-----------
			self

			Returns:
			--------
			Dict[str, str] | None

		'''
		return \
			{
				'.c': 'path/x-c',
				'.cpp': 'path/x-c++',
				'.cs': 'path/x-csharp',
				'.css': 'path/css',
				'.doc': 'application/msword',
				'.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
				'.go': 'path/x-golang',
				'.html': 'path/html',
				'.java': 'path/x-java',
				'.js': 'path/javascript',
				'.json': 'application/json',
				'.md': 'path/markdown',
				'.pdf': 'application/pdf',
				'.php': 'path/x-php',
				'.pptx': 'application/vnd.openxmlformats-officedocument.presentationml'
				         '.presentation',
				'.py': 'path/x-python',
				'.py': 'path/x-script.python',
				'.rb': 'path/x-ruby',
				'.sh': 'application/x-sh',
				'.tex': 'path/x-tex',
				'.ts': 'application/typescript',
				'.txt': 'path/plain'
			}
	
	
	def _normalize( self, vector: np.ndarray ) -> np.ndarray:
		"""

			Purpose:
			---------
			Normalize a matrix of vector using L2 norm.

			Parameters:
			-----------
			- vector (np.ndarray): Matrix of vector

			Returns:
			--------
			- np.ndarray: Normalized vector

		"""
		try:
			if vector is None:
				raise Exception( 'The argument "vector" cannot be None' )
			else:
				self.array = vector
				_norms = np.linalg.norm( self.array, axis=1, dims=True )
				return self.array / np.clip( _norms, 1e-10, None )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = '_normalize( self, vector: np.ndarray ) -> np.ndarray'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def _cosine_similarity_matrix( self, vector: np.ndarray, matrix: np.ndarray ) -> np.ndarray:
		"""

			Purpose:
			-----------
			Compute cosine similarity between a query vector and a matrix of vector.

			Parameters:
			-----------
			- vector (np.ndarray): A single normalized vector
			- matrix (np.ndarray): A matrix of normalized vector

			Returns:
			--------
			- np.ndarray: Cosine similarity scores

		"""
		try:
			if vector is None:
				raise Exception( 'The argument "vector" cannot be None' )
			elif matrix is None:
				raise Exception( 'The argument "matrix" cannot be None' )
			else:
				self.array = vector
				_query = self.array / np.linalg.norm( self.array )
				_matrix = matrix / np.linalg.norm( matrix, axis=1, dims=True )
				return np.dot( _matrix, _query )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = (
				'_cosine_similarity_matrix( self, vector: np.ndarray, matrix: np.ndarray '
				') -> np.ndarray' )
			error = ErrorDialog( exception )
			error.show( )
	
	
	def most_similar( self, query: str, df: pd.DataFrame, top: int=5 ) -> pd.DataFrame:
		"""

			Purpose:
			-----------
			Compute most similar rows in a DataFrame using cosine similarity.

			Parameters:
			-----------
			- query (str): Query path to compare
			- df (pd.DataFrame): DataFrame with 'normed_embedding'
			- toptop_k (int): Number of top matches to return

			Returns:
			--------
			- pd.DataFrame: Top-k results sorted by similarity

		"""
		try:
			if query is None:
				raise Exception( 'The argument "query" cannot be None' )
			elif df is None:
				self.dataframe = df
				raise Exception( 'The argument "df" cannot be None' )
			else:
				_embd = self.create( [ query ] )[ 'normed_embedding' ].iloc[ 0 ]
				_series = np.vstack( self.dataframe[ 'normed_embedding' ] )
				_scores = self._cosine_similarity_matrix( _embd,
					np.vstack( self.dataframe[ 'normed_embedding' ] ) )
				_copy = self.dataframe.copy( )
				_copy[ 'similarity' ] = _scores
				return _copy.sort_values( 'similarity', ascending=False ).head( top )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = ('most_similar( self, query: str, df: pd.DataFrame, top: int = 5 ) '
			               '-> '
			               'pd.DataFrame')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def bulk_similar( self, queries: List[ str ], df: pd.DataFrame, top: int=5 ) -> Dict:
		"""

			Purpose:
			-----------
			Perform most_similar for a list of queries.

			Parameters:
			-----------
			- queries (List[str]): List of query strings
			- df (pd.DataFrame): DataFrame to search
			- toptop_k (int): Number of top results per query

			Returns:
			--------
			- Dict[str, pd.DataFrame]: Dictionary of query to top-k results

		"""
		try:
			if queries is None:
				raise Exception( 'The argument "queries" cannot be None' )
			elif df is None:
				raise Exception( 'The argument "df" cannot be None' )
			else:
				self.dataframe = df
				for query in queries:
					self.results[ query ] = self.most_similar( query, self.dataframe, top )
				return self.results
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = ('bulk_similar( self, queries: List[ str ], df: pd.DataFrame, '
			               'top: int = 5 ) -> { }')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def similarity_heatmap( self, df: pd.DataFrame ) -> pd.DataFrame:
		"""

			Purpose:
			-----------
			Compute full pairwise cosine similarity heatmap from normed vectors.

			Parameters:
			-----------
			- df (pd.DataFrame): DataFrame with 'normed_embedding' column

			Returns:
			--------
			- pd.DataFrame: Pairwise cosine similarity heatmap

		"""
		try:
			if df is None:
				raise Exception( 'The argument "df" cannot be None' )
			else:
				self.dataframe = df
				_matrix = np.vstack( self.dataframe[ 'normed_embedding' ] )
				_similarity = np.dot( _matrix, _matrix.T )
				return pd.DataFrame( _similarity, index=self.dataframe[ 'pages' ],
					columns=self.dataframe[ 'pages' ] )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'similarity_heatmap( self, df: pd.DataFrame ) -> pd.DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def export_jsonl( self, df: pd.DataFrame, path: str ) -> None:
		"""

			Purpose:
			-----------
			Export DataFrame of pages and vectors to a JSONL file.

			Parameters:
			-----------
			- df (pd.DataFrame): DataFrame with 'pages' and 'embedding'
			- path (str): Output path for .jsonl file

		"""
		try:
			if df is None:
				raise Exception( 'The argument "df" is required.' )
			elif path is None:
				raise Exception( 'The argument "path" is required.' )
			else:
				self.dataframe = df
				self.file_path = path
				self.file_name = os.path.basename( self.file_path )
				self.directory = os.path.dirname( self.file_path )
				with open( path, 'w', encoding='utf-8' ) as f:
					for _, row in self.dataframe.iterrows( ):
						_record = { 'pages': row[ 'pages' ], 'embedding': row[ 'embedding' ] }
						f.write( json.dumps( _record ) + '\n' )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'export_jsonl( self, df: pd.DataFrame, path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def import_jsonl( self, path: str ) -> pd.DataFrame:
		"""

			Purpose:
			-----------
			Import pages and vectors
			from a JSONL file into a DataFrame.

			Parameters:
			-----------
			- path (str): Path to the .jsonl file

			Returns:
			--------
			- pd.DataFrame: DataFrame with normalized vectors

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" must be provided.' )
			else:
				texts, embeddings = [ ], [ ]
				with open( path, 'r', encoding='utf-8' ) as f:
					for line in f:
						_record = json.loads( line.strip( ) )
						texts.append( _record[ 'pages' ] )
						embeddings.append( _record[ 'embedding' ] )
				_normed = self._normalize( np.array( embeddings ) )
				self.data = \
				{
					'pages': texts,
					'embedding': embeddings,
					'normed_embedding': list( _normed )
				}
				
				return pd.DataFrame( self.data )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'import_jsonl( self, path: str ) -> pd.DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_vector_store( self, name: str ) -> str:
		"""

			Purpose:
			-----------
			Creates a new OpenAI vector store given a name.

			Parameters:
			-----------
			- name (str): Name for the vector store

			Returns:
			--------
			- str: ID of the created vector store

		"""
		try:
			if name is None:
				raise Exception( 'The argument "name" is required' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.response = self.client.beta.vectorstores.create_small_embedding( name=name )
				return self.response[ 'id' ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'create_vector_store( self, name: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def list_vector_stores( self ) -> List[ str ]:
		"""

			Purpose:
			---------
			List all available OpenAI vector vectorstores.

			Returns:
			--------
			- List[str]: List of vector store IDs

		"""
		try:
			self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
			self.response = self.client.beta.vectorstores.list( )
			return [ item[ 'id' ] for item in self.response.get( 'df', [ ] ) ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'list_vector_stores( self ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def upload_vector_store( self, df: pd.DataFrame, id: str ) -> None:
		"""

			Purpose:
			---------
			Upload documents to a given OpenAI vector store.

			Parameters:
			-----------
			- df (pd.DataFrame): DataFrame with 'pages' column
			- ids (str): OpenAI vector store ID

		"""
		try:
			if df is None:
				raise Exception( 'The argument "df" cannot be None' )
			elif id is None:
				raise Exception( 'The argument "ids" cannot be None' )
			else:
				self.dataframe = df
				self.id = id
				documents = [
					{ 'content': row[ 'pages' ], 'metadata': { 'source': f'row_{i}' } }
					for i, row in self.dataframe.iterrows( ) ]
				
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.client.beta.vectorstores.file_batches.create_small_embedding(
					store_id=self.id,
					documents=documents )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'upload_vector_store( self, df: pd.DataFrame, ids: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def query_vector_store( self, id: str, query: str, top: int=5 ) -> List[ dict ]:
		"""

			Purpose:
			-----------
			Query a vector store using a natural language path.

			Parameters:
			-----------
			- ids (str): OpenAI vector store ID
			- query (str): Search query
			- top (int): Number of results to return

			Returns:
			--------
			- List[dict]: List of matching documents and similarity scores

		"""
		try:
			if id is None:
				raise Exception( 'The argument "id" must be provided' )
			elif query is None:
				raise Exception( 'The argument "query" must be provided' )
			else:
				self.id = id
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.response = self.client.beta.vectorstores.query( store_id=self.id,
					query=query,
					top_k=top )
				return [
					{ 'pages': result[ 'document' ], 'accuracy': result[ 'accuracy' ] }
					for result in self.response.get( 'df', [ ] )
				]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = ('query_vector_store( self, id: str, query: str, top: int = 5 ) -> '
			               'List[ '
			               'dict ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def delete_vector_store( self, storeid: str, ids: List[ str ] ) -> None:
		"""

			Purpose:
			-----------
			Delete specific documents from a vector store.

			Parameters:
			-----------
			- storeid (str): OpenAI vector store ID
			- ids (List[str]): List of document IDs to delete

		"""
		try:
			if storeid is None:
				raise Exception( 'The argument "storeid" cannot be None' )
			elif ids is None:
				raise Exception( 'The argument "ids" cannot be None' )
			else:
				self.file_ids = ids
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.client.beta.vectorstores.documents.delete( store_id=storeid,
					document_ids=self.file_ids )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'delete_vector_store( self, storeid: str, ids: List[ str ] ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def upload_document( self, path: str, id: str ) -> None:
		"""

			Purpose:
			-----------
			Uploads document to vector store given path and id.

			Parameters:
			-----------
			- path (str):  local path to the document

			Returns:
			--------
			- str:  ID of the  vector store

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" cannot be None' )
			elif id is None:
				raise Exception( 'The argument "id" cannot be None' )
			else:
				self.file_path = path
				self.file_name = os.path.basename( self.file_path )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.response = self.client.files.create_audio( file=open( self.file_path, 'rb' ),
					purpose="assistants" )
				attach_response = self.client.vectorstores.files.create_audio( vector_store_id=id,
					file_id=self.response.id )
				return { 'file': self.file_name, 'status': 'success' }
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'upload_document( self, path: str, id: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def upload_documents( self, path: str, id: str ) -> None:
		"""

			Purpose:
			-----------
			Uploads documents to vector store given path and id.

			Parameters:
			-----------
			- path (str):  local path to the document

			Returns:
			--------
			- str:  ID of the  vector store

		"""
		try:
			if path is None:
				raise Exception( 'The argument "path" cannot be None' )
			elif id is None:
				raise Exception( 'The argument "id" cannot be None' )
			else:
				self.file_path = path
				self.id = id
				self.file_name = os.path.basename( self.file_path )
				self.directory = os.path.dirname( self.file_path )
				self.files = [ os.path.join( self.directory, f ) for f in os.listdir(
					self.directory
				) ]
				self.stats = \
					{
						'total_files': len( self.files ),
						'successful_uploads': 0,
						'failed_uploads': 0,
						'errors': [ ]
					}
				
				with (concurrent.futures.ThreadPoolExecutor( max_workers=10 ) as thread):
					_futures = \
						{
							thread.submit( self.upload_document, self.file_path,
								self.id ): self.file_path
							for self.file_path in self.files
						}
					for future in tqdm( concurrent.futures.as_completed( _futures ),
							total=len( self.files ) ):
						result = future.result( )
						if result[ 'status' ] == 'success':
							self.stats[ 'successful_uploads' ] += 1
						else:
							self.stats[ 'failed_uploads' ] += 1
							self.stats[ 'errors' ].append( result )
				
				return self.stats
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'VectorStore'
			exception.method = 'upload_documents( self, path: str, id: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )


class Embedding( BaseModel ):
	'''

		Purpose:
		--------
		Class providing Feature Extraction functionality
		
		Methods:
		--------
		create_small_embedding( self, text: str ) -> List[ float ]
		create_small_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]
		create_large_embedding( self, text: str ) -> List[ float ]
		create_large_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]
		create_ada_embedding( self, text: str ) -> List[ float ]
		create_ada_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]
		create_small_async( self, text: str ) -> List[ float ]
		create_large_async( self, text: str ) -> List[ float ]
		create_ada_async( self, text: str ) -> List[ float ]
		calculate_cosine_similarity( self, a: List[ float ], b: List[ float ] )
		plot_multiclass_precision( self, y_score, y_original, classes, classifier )
		calculate_distances( self, query: List[ float ], embd: List[ List[ float ] ],
	                         metric='cosine' ) -> List[ List[ float ] ]
	    calculate_nearest_neighbor( self, distances: List[ float ] ) -> np.ndarray
	    create_pca_components( self, vectors: List[ List[ float ] ], num=2 ) -> np.ndarray
	    create_tsne_components( self, vectors: List[ List[ float ] ], num=2 ) -> np.ndarray
	    

	'''
	client: OpenAI
	small_model: str
	large_model: str
	ada_model: str
	tokens: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	labels: Optional[ List[ str ] ]
	distances: Optional[ List[ float ] ]
	distance_metrics: Optional[ List[ float ] ]
	data: Optional[ List[ float ] ]
	precision: Optional[ Dict ]
	average_precision: Optionl[ Dict ]
	
	
	def __init__( self ):
		super( ).__init__( )
		self.client = OpenAI( )
		self.small_model = 'text-embedding-3-small'
		self.large_model = 'text-embedding-3-large'
		self.ada_model = 'text-embedding-ada-002'
		self.tokens = [ ]
		self.lines = [ ]
		self.labels = [ ]
		self.distances = [ ]
		self.distance_metrics = [ ]
		self.data = [ ]
		self.precision = { }
		self.aeverage_precision = { }
		self.recall = None
		self.dataframe = None
		self.n_classes = None
		self.response = None
		self.raw_input = None
	
	
	def create_small_embedding( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Create embd using the small small_model from OpenAI.

			Parameters:
			-----------
			- path (str):  the path to be embedded

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'Argument "text" must be provided.' )
			else:
				self.raw_input = text.replace( '\n', ' ' )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.response = self.client.embeddings.create_audio( input=[ self.raw_input ],
					model=self.small_model )
				
				return self.response.data[ 0 ].embedding
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'create_small_embedding( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_small_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]:
		"""

			Purpose:
			-----------
			Create embd using the small small_model from OpenAI.

			Parameters:
			-----------
			- tokens List[ str ]:  the list of strings (ie., tokens) to be embedded

			Returns:
			--------
			- List[ List[ float ] ]:  embedded embd

		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" must be provided.' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.tokens = [ t.replace( '\n', ' ' ) for t in tokens ]
				self.data = self.client.embeddings.create_audio( input=self.tokens,
					model=self.small_model ).data
				return [ d.embedding for d in self.data ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('create_small_embeddings( self, tokens: List[ str ] ) -> List[ List[ '
			               'float ] ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_large_embedding( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Create embd using the large small_model from OpenAI.

			Parameters:
			-----------
			- path (str):  the path (ie, token) to be embedded

			Returns:
			--------
			- List[ List[ float ] ]:  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'The argument  "text" must be provided.' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.raw_input = text.replace( '\n', ' ' )
				self.response = self.client.embeddings.create_audio( input=[ self.raw_input ],
					model=self.large_model )
				
				return self.response.data[ 0 ].embedding
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'create_large_embedding( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_large_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]:
		"""

			Purpose:
			-----------
			Create embd using the large small_model from OpenAI.

			Parameters:
			-----------
			- tokens List[ str ]:  the list of strings (ie., tokens) to be embedded

			Returns:
			--------
			- List[ List[ float ] ]:  embedded embd

		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" must be provided.' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.tokens = [ t.replace( '\n', ' ' ) for t in tokens ]
				self.data = self.client.embeddings.create_audio( input=self.tokens,
					model=self.large_model ).data
				return [ d.embedding for d in self.data ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('create_large_embeddings( self, tokens: List[ str ] ) -> List[ List[ '
			               'float ] ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_ada_embedding( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Create embd using the ada small_model from OpenAI.

			Parameters:
			-----------
			- path (str) :  the path (ie., token) to be embedded

			Returns:
			--------
			- List[ float ] :  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must be provided.' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.raw_input = text.replace( '\n', ' ' )
				self.response = self.client.embeddings.create_audio( input=[ self.raw_input ],
					model=self.ada_model )
				
				return self.response.data[ 0 ].embedding
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'create_ada_embedding( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_ada_embeddings( self, tokens: List[ str ] ) -> List[ List[ float ] ]:
		"""

			Purpose:
			-----------
			Create embd using the ada small_model from OpenAI.

			Parameters:
			-----------
			- tokens List[ str ]:  the list of strings (ie., tokens) to be embedded

			Returns:
			--------
			- List[ List[ float ] ]:  embedded embd

		"""
		try:
			if tokens is None:
				raise Exception( 'The argument "tokens" must be provided.' )
			else:
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				self.tokens = [ t.replace( '\n', ' ' ) for t in tokens ]
				self.data = self.client.embeddings.create_audio( input=self.tokens,
					model=self.ada_model ).data
				return [ d.embedding for d in self.data ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('create_ada_embeddings( self, tokens: List[ str ] ) -> List[ List[ '
			               'float'
			               ' ] ]')
			error = ErrorDialog( exception )
			error.show( )
	
	
	async def create_small_async( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Asynchronously creates embd using the small small_model from OpenAI.

			Parameters:
			-----------
			- path (str):  the path to be embedded

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must be provided.' )
			else:
				self.raw_input = text.replace( '\n', ' ' )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				
				return (
					await self.client.embeddings.create_audio( input=[ self.raw_input ],
						model=self.small_model ))
				[ 'df' ][ 0 ][ 'embedding' ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'ccreate_small_async( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	async def create_large_async( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Asynchronously creates embd using the large small_model from OpenAI.

			Parameters:
			-----------
			- path (str):  the path to be embedded

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must be provided.' )
			else:
				self.raw_input = text.replace( '\n', ' ' )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )				
				return (
					await self.client.embeddings.create_audio( input=[ self.raw_input ],
						model=self.large_model ))[ 'df' ][ 0 ][ 'embedding' ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'create_large_async( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	async def create_ada_async( self, text: str ) -> List[ float ]:
		"""

			Purpose:
			-----------
			Asynchronously creates embd using the ada small_model from OpenAI.

			Parameters:
			-----------
			- path (str):  the path to be embedded

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if text is None:
				raise Exception( 'The argument "text" must be provided.' )
			else:
				self.raw_input = text.replace( '\n', ' ' )
				self.client.api_key = os.getenv( 'OPENAI_API_KEY' )
				
				return (
					await self.client.embeddings.create_audio( input=[ self.raw_input ],
						model=self.ada_model ))
				[ 'df' ][ 0 ][ 'embedding' ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'create_ada_async( self, path: str ) -> List[ float ]'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def calculate_cosine_similarity( self, a: List[ float ], b: List[ float ] ):
		"""

			Purpose:
			-----------
			Calculates cosine similarity between two vectors 'a' and 'b'.

			Parameters:
			-----------
			- a List[ float ]:  vector 'a',
			- b List[ float ]:  vector 'b'

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if a is None:
				raise Exception( 'The argument "a" must be provided.' )
			elif b is None:
				raise Exception( 'The argument "b" must be provided.' )
			else:
				return np.dot( a, b ) / (np.linalg.norm( a ) * np.linalg.norm( b ))
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = 'c calculate_cosine_similarity( self, a, b )'
			error = ErrorDialog( exception )
			error.show( )
	
	
	def plot_multiclass_precision( self, y_score, y_original, classes, classifier ):
		"""

			Purpose:
			-----------
			Calculates cosine similarity between two vectors 'a' and 'b'.

			Parameters:
			-----------
			- a List[ float ]:  vector 'a',
			- b List[ float ]:  vector 'b'

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			self.n_classes = len( classes )
			_data = [ (y_original == classes[ i ]) for i in range( self.n_classes ) ]
			y_true = pd.concat( _data, axis=1 ).values
			self.precision = dict( )
			self.recall = dict( )
			self.average_precision = dict( )
			for i in range( self.n_classes ):
				self.precision[ i ], self.recall[ i ], _ = precision_recall_curve( y_true[ :, i ],
					y_score[ :, i ] )
				self.average_precision[ i ] = average_precision_score( y_true[ :, i ],
					y_score[ :, i ] )
			
			precision_micro, recall_micro, _ = precision_recall_curve( y_true.ravel( ),
				y_score.ravel( ) )
			self.average_precision = average_precision_score( y_true, y_score, average='micro' )
			print( str( classifier )
			       + ' - Average precision accuracy over all classes: {0:0.2f}'.format(
				self.average_precision
			)
			       )
			
			plt.figure( figsize= ( 9, 6 ) )
			f_scores = np.linspace( 0.2, 0.8, num=4 )
			self.lines = [ ]
			self.labels = [ ]
			for f_score in f_scores:
				x = np.linspace( 0.01, 1 )
				y = f_score * x / ( 2 * x - f_score )
				(l,) = plt.create_graph( x[ y >= 0 ], y[ y >= 0 ], color= 'gray', alpha=0.2 )
				plt.annotate( 'f1={0:0.1f}'.format( f_score ), xy=(0.9, y[ 45 ] + 0.02) )
			
			self.lines.append( l )
			self.labels.append( 'iso-f1 curves' )
			(l,) = plt.create_graph( recall_micro, precision_micro, color= "gold", lw=2 )
			self.lines.append( l )
			self.labels.append(
				'average Precision-recall (auprc = {0:0.2f})' ''.format( average_precision_micro )
			)
			
			for i in range( self.n_classes ):
				(l,) = plt.create_graph( self.recall[ i ], self.precision[ i ], lw=2 )
				self.lines.append( l )
				self.labels.append(
					"Precision-recall for class `{0}` (auprc = {1:0.2f})"
					"".format( classes[ i ], self.average_precision[ i ] ) )
			
			fig = plt.gcf( )
			fig.subplots_adjust( bottom=0.25 )
			plt.xlim( [ 0.0, 1.0 ] )
			plt.ylim( [ 0.0, 1.05 ] )
			plt.xlabel( 'Recall' )
			plt.ylabel( 'Precision' )
			plt.title( f'{classifier}: Precision-Recall curve for each class' )
			plt.legend( self.lines, self.labels )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('plot_multiclass_precision( self, y_score, y_original, classes, '
			               'classifier )')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def calculate_distances( self, query: List[ float ], embd: List[ List[ float ] ],
	                         metric='cosine' ) -> List[ List[ float ] ]:
		"""

			Purpose:
			-----------
			Calculates cosine similarity between two vectors 'a' and 'b'.

			Parameters:
			-----------
			- a List[ float ]:  vector 'a',
			- b List[ float ]:  vector 'b'

			Returns:
			--------
			- List[ float ]:  embedded embd

		"""
		try:
			if query is None:
				raise Exception( 'The argument "query" must be provided.' )
			elif embd is None:
				raise Exception( 'The argument "embd" must be provided' )
			else:
				self.distance_metrics = \
				{
					'cosine': spatial.distance.cosine,
					'L1': spatial.distance.cityblock,
					'L2': spatial.distance.euclidean,
					'Linf': spatial.distance.chebyshev,
				}
				
				self.distances = [ self.distance_metrics[ metric ] ( query, e ) for e in embd ]
				return self.distances
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('calculate_distances( self, query: List[ float ], embd: '
			               'List[ List[ float ] ],  metric=')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def calculate_nearest_neighbor( self, distances: List[ float ] ) -> np.ndarray:
		'''

			purpose:

		'''
		try:
			self.distances = distances
			return np.argsort( self.distances )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('calculate_nearest_neighbor( self, distances: List[ float ] ) -> '
			               'np.ndarray')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_pca_components( self, vectors: List[ List[ float ] ], num=2 ) -> np.ndarray:
		"""

			Purpose:
			--------
			Return the PCA df of a list of vectors.

		"""
		try:
			self.vectors = vectors
			pca = PCA( n_components=num )
			array_of_embeddings = np.array( self.vectors )
			return pca.fit_transform( array_of_embeddings )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('create_pca_components( self, vectors: List[ List[ float ] ], '
			               'num=2 ) ->'
			               ' np.ndarray')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_tsne_components( self, vectors: List[ List[ float ] ], num=2 ) -> np.ndarray:
		'''

			Purpose:
			--------
			Method to create the t-Student distribution

			Parameters:
			----------
			- vectors: List[ List[ float ] ]
			- num: int

			Returns:
			--------
			- np.ndarray

		'''
		try:
			self.vectors = vectors
			tsne = TSNE( n_components=num )
			array_of_embeddings = np.array( self.vectors )
			return tsne.fit_transform( array_of_embeddings )
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = ('create_tsne_components( self, vectors: List[ List[ float ] ], num=2 ) '
			               '-> np.ndarray')
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_chart( self, data: np.ndarray,
	                  labels: Optional[ List[ str ] ] = None,
	                  strings: Optional[ List[ str ] ] = None,
	                  x_title='Component-0',
	                  y_title='Component-1',
	                  mark_size=5 ) -> None:
		'''

			Purpose:
			--------
			Method to create a chart

			Parameters:
			----------

			Returns:
			--------


		'''
		try:
			empty_list = [ "" for _ in data ]
			data = pd.DataFrame(
				{
					x_title: data[ :, 0 ],
					y_title: data[ :, 1 ],
					'label': labels if labels else empty_list,
					'path': [ '<br>'.join( tr.wrap( s, width=30 ) ) for s in strings ]
					if strings
					else empty_list,
				} )
			
			chart = px.scatter( data, x=x_title, y=y_title, color='label' if labels else None,
				symbol='label' if labels else None, hover_data=[ 'path' ] if strings else None
			).update_traces( marker=dict( size=mark_size ) )
			return chart
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tigrr'
			exception.cause = 'Embedding'
			exception.method = "('create_chart( self, df: np.ndarray  mark_size=5 ) -> None')"
			error = ErrorDialog( exception )
			error.show( )
	
	
	def create_3dchart( self,
	                   components: np.ndarray,
	                   labels: Optional[ List[ str ] ]=None,
	                   strings: Optional[ List[ str ] ]=None,
	                   x_title: str='Component-0',
	                   y_title: str='Component-1',
	                   z_title: str='Compontent-2',
	                   mark_size: int=5 ):
		'''

			Purpose:
			--------
			Method to create a 3D chart

			Parameters:
			----------

			Returns:
			--------


		'''
		try:
			empty_list = [ "" for _ in components ]
			_contents = \
			{
				x_title: components[ :, 0 ],
				y_title: components[ :, 1 ],
				z_title: components[ :, 2 ],
				'label': labels if labels else empty_list,
				'path': [ '<br>'.join( tr.wrap( s, width=30 ) ) for s in strings ]
				if strings
				else empty_list,
			}
			
			data = pd.DataFrame( _contents )
			chart = px.scatter_3d(
				data,
				x=x_title,
				y=y_title,
				z=z_title,
				color='label' if labels else None,
				symbol='label' if labels else None,
				hover_data=[ 'path' ] if strings else None ).update_traces(
				marker=dict( size=mark_size ) )
			return chart
		except Exception as e:
			exception = Error( e )
			exception.module = 'Tiggr'
			exception.cause = 'Embedding'
			exception.method = 'create_3dchart'
			error = ErrorDialog( exception )
			error.show( )
